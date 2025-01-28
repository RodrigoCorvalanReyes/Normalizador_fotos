import os
import tkinter as tk
from tkinter import filedialog, ttk, messagebox

# -----------
# Para imágenes
# -----------
from PIL import Image

# -----------
# Para DOCX
# -----------
from docx import Document
from docx.shared import Inches


# --------------------------------------------------------------------
# Funciones de lógica (obtener estructura, procesar imágenes,
# generar DOCX).
# --------------------------------------------------------------------

def obtener_estructura_imagenes(ruta_base):
    """
    Recorre la carpeta 'ruta_base' recursivamente y devuelve una lista
    de diccionarios con esta forma:
        [
            {
                'ruta': 'C:/.../carpeta',
                'imagenes': ['foto1.jpg', 'foto2.png', ...]
            },
            ...
        ]
    """
    estructura = []
    for root, dirs, files in os.walk(ruta_base):
        # Filtra solo archivos de imagen (agrega extensiones si necesitas)
        imagenes = [f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
        if imagenes:
            estructura.append({
                'ruta': root,
                'imagenes': imagenes
            })
    return estructura


def procesar_imagen(ruta_imagen, ruta_salida, ancho=1280, alto=720, calidad=85):
    """
    Abre la imagen en 'ruta_imagen', la convierte a JPG, la redimensiona a (ancho x alto)
    y la guarda en 'ruta_salida' con la calidad especificada.
    """
    with Image.open(ruta_imagen) as img:
        # Convertir a RGB (por si es PNG con canal alpha)
        img = img.convert('RGB')
        # Redimensionar (forzado). Para mantener proporciones, haz un cálculo previo.
        img = img.resize((ancho, alto))
        
        # Crear carpeta de salida si no existe
        os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
        # Guardar como JPG con la calidad especificada
        img.save(ruta_salida, 'JPEG', quality=calidad)


def construir_ruta_salida(ruta_base, ruta_destino, root, nombre_imagen):
    """
    Construye la ruta de salida para la imagen, manteniendo la estructura de carpetas
    con base en la ruta_base y el directorio actual 'root'.
    """
    ruta_relativa = os.path.relpath(root, ruta_base)
    return os.path.join(ruta_destino, ruta_relativa, nombre_imagen)


def generar_docx(estructura, ruta_base, ruta_salida, ruta_docx):
    """
    Genera un archivo DOCX en 'ruta_docx' donde cada sección corresponde
    a la jerarquía de carpetas encontrada, e incluye las imágenes procesadas.
    """
    document = Document()
    document.add_heading('Documento de Imágenes', level=1)

    for elemento in estructura:
        # Calculamos la ruta relativa y usamos cada carpeta como parte del título
        ruta_carpetas = os.path.relpath(elemento['ruta'], ruta_base)
        secciones = ruta_carpetas.split(os.sep)
        titulo_seccion = ", ".join([s.replace("_", " ").title() for s in secciones])
        
        # Insertamos un encabezado (nivel 2) con el nombre de la sección
        document.add_heading(titulo_seccion, level=2)

        # Para cada imagen en la carpeta
        for img_name in elemento['imagenes']:
            # La imagen procesada se encuentra en 'ruta_salida' con la misma ruta relativa
            ruta_imagen_procesada = construir_ruta_salida(ruta_base, ruta_salida, elemento['ruta'], img_name)
            
            # Insertamos la imagen en el DOCX
            document.add_picture(ruta_imagen_procesada, width=Inches(4.0))
            # Añade un párrafo debajo con el nombre
            document.add_paragraph(f"Imagen: {img_name}")

        # Espacio adicional entre secciones (opcional)
        document.add_paragraph("")

    document.save(ruta_docx)
    print(f"Documento DOCX generado en: {ruta_docx}")


def main(ruta_entrada, ruta_salida, ruta_docx, ancho=1280, alto=720, calidad=85):
    """
    Función principal:
      1) Obtiene la estructura de imágenes.
      2) Procesa/convierte cada imagen y la guarda en 'ruta_salida'.
      3) Genera un DOCX con la estructura y las imágenes procesadas.
    """
    # 1. Obtener estructura de imágenes
    estructura = obtener_estructura_imagenes(ruta_entrada)

    # 2. Procesar cada imagen
    for elem in estructura:
        root = elem['ruta']
        for img_name in elem['imagenes']:
            ruta_original = os.path.join(root, img_name)
            ruta_nueva = construir_ruta_salida(ruta_entrada, ruta_salida, root, img_name)
            procesar_imagen(ruta_original, ruta_nueva, ancho=ancho, alto=alto, calidad=calidad)

    # 3. Generar el documento DOCX
    generar_docx(estructura, ruta_entrada, ruta_salida, ruta_docx)


# --------------------------------------------------------------------
# GUI con tkinter para que el usuario seleccione directorios,
# configure parámetros y genere el DOCX con un clic.
# --------------------------------------------------------------------

class Aplicacion:
    def __init__(self, root):
        self.root = root
        self.root.title("RCSoftware")

        # Variables de control con valores por defecto
        self.ruta_entrada = tk.StringVar()
        self.ruta_salida = tk.StringVar()
        self.ruta_docx = tk.StringVar()
        # Por defecto 1280x720
        self.ancho = tk.IntVar(value=1280)
        self.alto = tk.IntVar(value=720)
        # Calidad por defecto 85
        self.calidad = tk.IntVar(value=85)

        # Diseño de la interfaz
        ttk.Label(root, text="Directorio de entrada:").grid(row=0, column=0, padx=5, pady=5, sticky='w')
        ttk.Entry(root, textvariable=self.ruta_entrada, width=40).grid(row=0, column=1, padx=5, pady=5, sticky='w')
        ttk.Button(root, text="Buscar", command=self.seleccionar_directorio_entrada).grid(row=0, column=2, padx=5, pady=5)

        ttk.Label(root, text="Directorio de salida:").grid(row=1, column=0, padx=5, pady=5, sticky='w')
        ttk.Entry(root, textvariable=self.ruta_salida, width=40).grid(row=1, column=1, padx=5, pady=5, sticky='w')
        ttk.Button(root, text="Buscar", command=self.seleccionar_directorio_salida).grid(row=1, column=2, padx=5, pady=5)

        ttk.Label(root, text="Archivo DOCX:").grid(row=2, column=0, padx=5, pady=5, sticky='w')
        ttk.Entry(root, textvariable=self.ruta_docx, width=40).grid(row=2, column=1, padx=5, pady=5, sticky='w')
        ttk.Button(root, text="Guardar como...", command=self.seleccionar_archivo_docx).grid(row=2, column=2, padx=5, pady=5)

        ttk.Label(root, text="Ancho (px):").grid(row=3, column=0, padx=5, pady=5, sticky='w')
        ttk.Entry(root, textvariable=self.ancho, width=10).grid(row=3, column=1, sticky='w', padx=5, pady=5)

        ttk.Label(root, text="Alto (px):").grid(row=4, column=0, padx=5, pady=5, sticky='w')
        ttk.Entry(root, textvariable=self.alto, width=10).grid(row=4, column=1, sticky='w', padx=5, pady=5)

        ttk.Label(root, text="Calidad JPG (1-95):").grid(row=5, column=0, padx=5, pady=5, sticky='w')
        ttk.Entry(root, textvariable=self.calidad, width=10).grid(row=5, column=1, sticky='w', padx=5, pady=5)

        ttk.Button(root, text="Procesar", command=self.procesar).grid(row=6, column=0, columnspan=3, pady=10)

    def seleccionar_directorio_entrada(self):
        carpeta = filedialog.askdirectory(title="Seleccionar carpeta de entrada")
        if carpeta:
            self.ruta_entrada.set(carpeta)

    def seleccionar_directorio_salida(self):
        carpeta = filedialog.askdirectory(title="Seleccionar carpeta de salida")
        if carpeta:
            self.ruta_salida.set(carpeta)

    def seleccionar_archivo_docx(self):
        archivo = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Guardar DOCX como..."
        )
        if archivo:
            self.ruta_docx.set(archivo)

    def procesar(self):
        """
        Llama a la función principal con los parámetros
        ingresados por el usuario en la interfaz.
        """
        ruta_entrada = self.ruta_entrada.get()
        ruta_salida = self.ruta_salida.get()
        ruta_docx = self.ruta_docx.get()
        ancho = self.ancho.get()
        alto = self.alto.get()
        calidad = self.calidad.get()

        # Verifica que no falten parámetros
        if not (ruta_entrada and ruta_salida and ruta_docx):
            messagebox.showwarning("Atención", "Por favor, selecciona todas las rutas.")
            return

        # Ejecuta el proceso
        try:
            main(ruta_entrada, ruta_salida, ruta_docx, ancho=ancho, alto=alto, calidad=calidad)
            # Muestra un mensaje indicando que terminó
            messagebox.showinfo("Proceso Completado", f"El documento se ha generado en:\n{ruta_docx}")
        except Exception as e:
            # Mensaje de error
            messagebox.showerror("Error", f"Ocurrió un error durante el procesamiento:\n{e}")


def main_gui():
    root = tk.Tk()
    app = Aplicacion(root)
    root.mainloop()


if __name__ == "__main__":
    # Si quieres usar la interfaz GUI, descomenta esto:
    main_gui()

    # Si prefieres usarlo en modo consola:
    # ruta_entrada = "C:/imagenes"
    # ruta_salida = "C:/salida"
    # ruta_docx = "C:/salida/DocumentoFinal.docx"
    # main(ruta_entrada, ruta_salida, ruta_docx, ancho=1280, alto=720, calidad=85)
    # print("Proceso completado.")
